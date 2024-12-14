from docx import Document
from docx.shared import Inches
import io
from openpyxl import Workbook
from openpyxl.drawing.image import Image
from PIL import Image as PILImage

#the things we pasted in word are mostly formed with hyperlinks, while if we paste without form(仅保留文本粘贴), we will lose most pictures
#so we must run another script to clean all hyperlinks.
def split_text_to_word_and_images_to_excel(input_path, text_output_path, excel_output_path):
    original_doc = Document(input_path)
    text_doc = Document()
    wb = Workbook()
    ws = wb.active
    ws.title = "Images"
    ws.append(["Image Number", "Image"]) 
    # 创建一个集合来跟踪已插入的图片，避免重复
    inserted_images = set()
    image_counter = 1
    for paragraph in original_doc.paragraphs:
        # 提取段落文字并存入文字文档
        if paragraph.text.strip():
            text_doc.add_paragraph(paragraph.text)

        # 提取段落中的图片并存入 Excel
        for run in paragraph.runs:
            if "graphic" in run._element.xml:  # 检查 run 中是否包含图片
                for rel in original_doc.part.rels.values():
                    if "image" in rel.target_ref:
                        if rel.target_ref not in inserted_images:
                            try:
                                image_data = rel.target_part.blob
                                
                                # 使用 PILImage 处理图片并插入 Excel
                                image_stream = io.BytesIO(image_data)
                                pil_image = PILImage.open(image_stream)
                                excel_image = Image(image_stream)
                                excel_image.width = 150  # 调整图片宽度（像素）
                                excel_image.height = int(pil_image.size[1] * (150 / pil_image.size[0]))  # 按比例调整高度

                                # 插入图片到 Excel
                                ws.add_image(excel_image, f"B{image_counter + 1}")  # B列用于存储图片
                                ws.cell(row=image_counter + 1, column=1, value=image_counter)  # A列存储图片编号

                                # 增加图片计数
                                image_counter += 1

                                # 标记图片已插入
                                inserted_images.add(rel.target_ref)
                            except KeyError:
                                print(f"无法找到图片: {rel.target_ref}")

    text_doc.save(text_output_path)
    print(f"文字内容已保存到 {text_output_path}")

    wb.save(excel_output_path)
    print(f"图片内容已保存到 {excel_output_path}")


# 文件路径
input_path = "D:\\5507\\after2.docx"  # 输入路径
text_output_path = "D:\\5507\\data3_text.docx"  # 文字路径
excel_output_path = "D:\\5507\\data3_images.xlsx"  # 图片路径

# 运行函数
split_text_to_word_and_images_to_excel(input_path, text_output_path, excel_output_path)
